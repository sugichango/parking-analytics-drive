import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_font(run, font_name='HG丸ｺﾞｼｯｸM-PRO'):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_spacing(paragraph, line_spacing=1.15, space_before=0, space_after=4):
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.line_spacing = line_spacing

def add_formatted_text(paragraph, text, font_name='HG丸ｺﾞｼｯｸM-PRO', font_size=10.5, default_bold=False):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run_text = part[2:-2]
            bold = True
        else:
            run_text = part
            bold = default_bold
        
        if not run_text:
            continue
            
        run = paragraph.add_run(run_text)
        run.bold = bold
        run.font.size = Pt(font_size)
        set_font(run, font_name)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120): # dxa単位
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # 0.5 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # ページ余白を標準（左右上下1.0インチ）に設定
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    in_quote = False
    quote_text = []
    
    font_name = 'HG丸ｺﾞｼｯｸM-PRO'
    
    def flush_table():
        nonlocal in_table, table_data
        if not table_data:
            return
        
        rows = len(table_data)
        cols = len(table_data[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)
        
        for r_idx, row_cells in enumerate(table_data):
            for c_idx, cell_text in enumerate(row_cells):
                if c_idx >= len(table.columns):
                    continue
                cell = table.cell(r_idx, c_idx)
                # セル内余白を少し拡大
                set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
                
                if r_idx == 0:
                    set_cell_background(cell, 'F2F2F2')
                    bold = True
                else:
                    bold = False
                
                p = cell.paragraphs[0]
                # テーブル内行間も少し広く設定
                set_paragraph_spacing(p, line_spacing=1.05, space_before=0, space_after=2)
                
                # 数値系データ（数字を含み、かつ数値・記号・単位で構成されている）かどうかの判定
                stripped_text = cell_text.replace('*', '').strip()
                has_digit = any(c.isdigit() for c in stripped_text)
                is_numeric = False
                if has_digit:
                    # 数字、カンマ、ドット、記号、特定の単位のみで構成されているかチェック
                    allowed_chars = re.match(r'^[\d,%\.\+\-▼▲万円台割\s→]+$', stripped_text)
                    if allowed_chars:
                        is_numeric = True

                if is_numeric:
                    p.alignment = 2 # 右寄せ
                else:
                    p.alignment = 0 # 左寄せ
                
                # プラス値が含まれる場合は太字にする
                cell_bold = bold
                if '+' in cell_text:
                    cell_bold = True
                    
                add_formatted_text(p, cell_text, font_name=font_name, font_size=9.5, default_bold=cell_bold)
                
        # テーブルの後にスペースを追加
        p_after = doc.add_paragraph()
        set_paragraph_spacing(p_after, line_spacing=1.15, space_before=0, space_after=6)
        
        table_data = []
        in_table = False

    def flush_quote():
        nonlocal in_quote, quote_text
        if not quote_text:
            return
        
        for qt in quote_text:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.15, space_before=2, space_after=4)
            
            # 箇条書きを含むかチェックしてぶら下げを設定
            bullet_sub_match = re.match(r'^\s*[-\*]\s+(.*)$', qt)
            if bullet_sub_match:
                content = bullet_sub_match.group(1)
                p.paragraph_format.left_indent = Pt(36) # 引用ブロックのインデント(Pt(22)) + 箇条書きインデント(Pt(14))
                p.paragraph_format.first_line_indent = Pt(-14)
                
                run_bullet = p.add_run('・')
                run_bullet.bold = True
                run_bullet.font.size = Pt(9.5)
                set_font(run_bullet, font_name)
                
                run_tab = p.add_run('\t')
                set_font(run_tab, font_name)
                
                add_formatted_text(p, content, font_name=font_name, font_size=9.5)
            else:
                p.paragraph_format.left_indent = Pt(22) # 約0.3インチ
                add_formatted_text(p, qt, font_name=font_name, font_size=10.0)
            
        quote_text = []
        in_quote = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        if not line.strip():
            if in_table:
                flush_table()
            if in_quote:
                flush_quote()
            i += 1
            continue
            
        if line.strip() == '---':
            if in_table: flush_table()
            if in_quote: flush_quote()
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.15, space_before=6, space_after=6)
            run = p.add_run('━' * 45)
            run.font.size = Pt(8)
            set_font(run, font_name)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue
            
        if line.startswith('>'):
            if in_table: flush_table()
            in_quote = True
            content = line[1:].strip()
            if content.startswith('-') or content.startswith('*'):
                content = '  ' + content
            quote_text.append(content)
            i += 1
            continue
        elif in_quote and not line.startswith('>'):
            flush_quote()
            
        if line.strip().startswith('|'):
            if in_quote: flush_quote()
            in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            is_sep = all(re.match(r'^:?-+:?$', c) for c in cells)
            if not is_sep:
                table_data.append(cells)
            i += 1
            continue
        elif in_table and not line.strip().startswith('|'):
            flush_table()
            
        if line.startswith('# '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.2, space_before=12, space_after=6)
            add_formatted_text(p, line[2:].strip(), font_name=font_name, font_size=16, default_bold=True)
            i += 1
            continue
            
        if line.startswith('## '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.2, space_before=10, space_after=4)
            add_formatted_text(p, line[3:].strip(), font_name=font_name, font_size=13, default_bold=True)
            i += 1
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, line_spacing=1.15, space_before=8, space_after=4)
            add_formatted_text(p, line[4:].strip(), font_name=font_name, font_size=11, default_bold=True)
            i += 1
            continue
            
        bullet_match = re.match(r'^(\s*)[-\*]\s+(.*)$', line)
        if bullet_match:
            indent_spaces = len(bullet_match.group(1))
            content = bullet_match.group(2)
            p = doc.add_paragraph()
            # ぶら下げインデントを設定 (1行目は左に飛び出し、2行目以降は揃う)
            level = indent_spaces // 2
            p.paragraph_format.left_indent = Pt(28 + 14 * level)
            p.paragraph_format.first_line_indent = Pt(-14)
            set_paragraph_spacing(p, line_spacing=1.15, space_before=0, space_after=3)
            
            run_bullet = p.add_run('・')
            run_bullet.bold = True
            run_bullet.font.size = Pt(10.5)
            set_font(run_bullet, font_name)
            
            # 記号の後にタブを入れて、ぶら下げ位置と本文開始位置を揃える
            run_tab = p.add_run('\t')
            set_font(run_tab, font_name)
            
            add_formatted_text(p, content, font_name=font_name, font_size=10.5)
            i += 1
            continue
            
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=1.15, space_before=0, space_after=4)
        add_formatted_text(p, line.strip(), font_name=font_name, font_size=10.5)
        i += 1

    if in_table: flush_table()
    if in_quote: flush_quote()
    
    saved_successfully = False
    current_docx_path = docx_path
    attempt = 1
    
    while not saved_successfully:
        try:
            doc.save(current_docx_path)
            print(f"Successfully saved to {current_docx_path}")
            saved_successfully = True
        except PermissionError:
            dir_name = os.path.dirname(docx_path)
            base_name = os.path.basename(docx_path)
            name, ext = os.path.splitext(base_name)
            if attempt == 1:
                current_docx_path = os.path.join(dir_name, f"{name}_調整版{ext}")
            else:
                current_docx_path = os.path.join(dir_name, f"{name}_調整版_{attempt}{ext}")
            attempt += 1
            if attempt > 10:
                print("Error: Could not save file, too many locked attempts.")
                break

if __name__ == '__main__':
    md_path = r'c:\Users\sugitamasahiko\Documents\parking_system\analysis\①_一般利用台数推移分析.md'
    docx_path = r'c:\Users\sugitamasahiko\Documents\parking_system\analysis\①_一般利用台数推移分析.docx'
    convert_md_to_docx(md_path, docx_path)
